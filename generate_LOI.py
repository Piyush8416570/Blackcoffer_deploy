

from ast import Not
import io
import os
from random import choice
from sys import exec_prefix
from warnings import catch_warnings
from git import Tree
import pandas as pd
from docx import Document
import streamlit as st
from datetime import date
from PyPDF2 import PdfReader
from pdf2docx import Converter
# from docx2pdf import convert
# import pythoncom
from sympy import true

# from streamlit.runtime.legacy_caching import cache_clear
# from streamlit.runtime.caching import cache_data
def convert_pdf_to_docx(pdf_path):
    docx_path = pdf_path.replace('.pdf', '.docx')
    cv = Converter(pdf_path)
    cv.convert(docx_path)
    cv.close()
    return docx_path

import subprocess

def convert_docx_to_pdf(docx_path):
    pdf_path = docx_path.replace('.docx', '.pdf')
    try:
        subprocess.run(['unoconv', '-f', 'pdf', '-o', pdf_path, docx_path], check=True)
        return pdf_path
    except subprocess.CalledProcessError as e:
        print(f"Error converting file: {e}")
        return None
# def convert_docx_to_pdf(docx_path):
#     # pythoncom.CoInitialize()  # Initialize COM library
#     pdf_path = docx_path.replace('.docx', '.pdf')
#     convert(docx_path, pdf_path)
#     # pythoncom.CoUninitialize()  # Uninitialize COM library after use
#     return pdf_path


def generate_lois(df, template_file, company_name, your_name):
    # Convert PDF to DOCX if necessary
    if template_file.name.lower().endswith('.pdf'):
        with open("temp_template.pdf", "wb") as f:
            if isinstance(template_file, io.BufferedReader):
                f.write(template_file.read())
            else:
                f.write(template_file.getvalue())
        template_path = convert_pdf_to_docx("temp_template.pdf")
    else:
        with open("temp_template.docx", "wb") as f:
            if isinstance(template_file, io.BufferedReader):
                f.write(template_file.read())
            else:
                f.write(template_file.getvalue())
        template_path = "temp_template.docx"

    # Read the template
    template_doc = Document(template_path)
    print(template_doc)

    # To check whether the Some data should be as same or user want to edit somethings
   
    generated_files = []
    for index, row in df.iterrows():
        try:
            data = row.to_dict()
            
            # Add today's date to the data dictionary
            data['Today Date'] = date.today().strftime("%d, %B, %Y")
    
    
            # Use the input values for all LOIs
            data['Company'] = company_name
            data['Your Name'] = your_name
    
            print(f'PDF is here:', (data))
            # Create a new document for each row
            doc = Document(template_path)
    
            # Replace placeholders
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    if f"{{{{{key}}}}}" in paragraph.text:
                        paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))
    
            # Save the document
            output_path = f"LOI_{index}.docx"
            doc.save(output_path)
            # now we are going to reconvert the docx to pdf 
            pdf_path = convert_docx_to_pdf(output_path)
    
            if os.path.exists(pdf_path):
                generated_files.append(pdf_path)
                st.success(f"Generated LOI for index {index}")
            else:
                st.warning(f"Failed to generate PDF for index {index}")
    
                # Optional: remove the intermediate DOCX file
            if os.path.exists(output_path):
                os.remove(output_path)

        except Exception as e:
            st.error(f"Error generating LOI for index {index}: {str(e)}")

    return generated_files

def handle_generate_lois(df):
    st.write("Generating LOIs...")
    
    # Initialize session state variables
    if 'use_default_template' not in st.session_state:
        st.session_state.use_default_template = False
    if 'uploaded_template' not in st.session_state:
        st.session_state.uploaded_template = None
    if 'company_name' not in st.session_state:
        st.session_state.company_name = "We Buy Houses Anywhere LLC"
    if 'your_name' not in st.session_state:
        st.session_state.your_name = "Justin Pickell"

    col1, col2 = st.columns([4, 1])

    with col1:
        uploaded_template = st.file_uploader("Upload LOI template (PDF or DOCX)", type=["pdf", "docx"])
        if uploaded_template is not None:
            st.session_state.uploaded_template = uploaded_template
            st.session_state.use_default_template = False

    with col2:
        use_default = st.checkbox("Use default template", value=st.session_state.use_default_template)
        st.session_state.use_default_template = use_default
        if use_default:
            if st.button("View default template"):
                with open("Generic_LOI.pdf", "rb") as file:
                    st.download_button(
                        label="Download default template",
                        data=file,
                        file_name="Generic_LOI.pdf",
                        mime="application/pdf"
                    )

    if st.session_state.uploaded_template is not None:
        st.write(f"Using uploaded template: {st.session_state.uploaded_template.name}")
    elif st.session_state.use_default_template:
        st.write("Using default template: Generic_LOI.pdf")
    else:
        st.warning("Please upload a template or use the default template to proceed.")
        return

    # Editable fields
    st.subheader("Edit LOI Information")
    company_name = st.text_input("Company Name", st.session_state.company_name)
    your_name = st.text_input("Your Name", st.session_state.your_name)

    # Update session state with new values
    st.session_state.company_name = company_name
    st.session_state.your_name = your_name

    if st.button("Generate LOIs"):
        if st.session_state.use_default_template:
            template = open("Generic_LOI.pdf", "rb")
        else:
            template = st.session_state.uploaded_template

        generated_files = generate_lois(df, template, company_name, your_name)
        st.success(f"Generated {len(generated_files)} LOIs.")
        
        st.download_button(
            label="Download All LOIs",
            data=generate_zip(generated_files),
            file_name="all_lois.zip",
            mime="application/zip"
        )

# Keep the generate_zip function as is
def generate_zip(file_paths):
    import io
    from zipfile import ZipFile

    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, "w") as zip_file:
        for file_path in file_paths:
            if file_path and os.path.exists(file_path):
                zip_file.write(file_path, os.path.basename(file_path))

    return zip_buffer.getvalue()

